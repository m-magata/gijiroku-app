import streamlit as st
def check_password():
    def password_entered():
        if st.session_state["password"] == st.secrets["APP_PASSWORD"]:
            st.session_state["password_correct"] = True
            del st.session_state["password"]
        else:
            st.session_state["password_correct"] = False

    if "password_correct" not in st.session_state:
        st.text_input("社内パスワードを入力してください", type="password", key="password", on_change=password_entered)
        return False
    elif not st.session_state["password_correct"]:
        st.text_input("社内パスワードを入力してください", type="password", key="password", on_change=password_entered)
        st.error("パスワードが違います")
        return False
    else:
        return True

if not check_password():
    st.stop()
import google.generativeai as pal_genai
from docx import Document
import io, os, re, time, tempfile
from datetime import datetime

# --- 1. ページ設定（UIのみ） ---
st.set_page_config(page_title="Audio Transcription", layout="wide")

# --- UI CSS（微調整：中央スペース／HowToオレンジ／番号位置／ステータス下げ） ---
st.markdown(
    """
<style>
/* タイトル切れ対策（維持） */
.block-container { padding-top: 2.8rem; padding-bottom: 1rem; }
div[data-testid="stVerticalBlock"] { gap: 0.65rem; }

/* タイトル */
.at-title{
  font-size: 44px;
  font-weight: 900;
  letter-spacing: 0.2px;
  line-height: 1.14;
  margin: 0.2rem 0 0.35rem 0;
}
.at-ver{
  font-size: 13px;
  color: #6b7280;
  font-weight: 800;
  margin-left: 10px;
  vertical-align: middle;
}

/* モデル表示（右上） */
.model-chip{
  text-align:right;
  padding: 8px 10px;
  border: 1px solid #bfdbfe;
  background: #eff6ff;
  border-radius: 12px;
  display: inline-block;
}
.model-chip .value{ font-size: 14px; font-weight: 900; color:#111827; }

/* 作成モード：wrap */
div[data-testid="stRadio"] [role="radiogroup"] {
  display:flex;
  flex-wrap:wrap;
  gap: 6px 16px;
}
div[data-testid="stRadio"] label[data-baseweb="radio"] span{ font-size: 13px; }

/* 入力を締める */
div[data-testid="stTextInput"] input { padding-top: 6px; padding-bottom: 6px; }
div[data-testid="stTextArea"] textarea { padding-top: 10px; padding-bottom: 10px; }

/* file_uploader：余白を詰める */
div[data-testid="stFileUploader"] section{
  padding-top: 6px !important;
  padding-bottom: 6px !important;
}
div[data-testid="stFileUploaderDropzone"]{
  padding-top: 10px !important;
  padding-bottom: 10px !important;
}

/* file_uploader 英語を「見た目上」消す */
div[data-testid="stFileUploaderDropzone"] small { display:none !important; }
div[data-testid="stFileUploaderDropzone"] p { display:none !important; }

/* 日本語案内 */
.upload-help{
  font-size: 12px;
  color:#6b7280;
  margin-top: 6px;
}

/* 追加指示の上余白 */
.section-spacer{ height: 10px; }

/* 生成ボタンを下げる */
.btn-spacer{ height: 16px; }

/* How to use（薄いオレンジ + 番号揃え） */
.howto-box{
  background: #fff7ed;         /* orange-50 */
  border: 1px solid #fed7aa;   /* orange-200 */
  border-radius: 12px;
  padding: 10px 12px;
  margin-top: 6px;
}
.howto-title{
  font-weight: 900;
  color:#111827;
  margin: 0 0 6px 0;
}

/* ここが「番号をHの下に」：余計なインデントを消す */
.howto{
  margin: 0;
  padding-left: 1.05em;        /* ← 微妙にだけ下げる（Hの開始位置に近づく） */
}
.howto li{
  margin: 2px 0;
  font-size: 12.5px;
}
.howto small{ color:#6b7280; }

/* ステータス（半分高さ） */
@keyframes blink { 0%{opacity:1;} 50%{opacity:0.35;} 100%{opacity:1;} }
.blinking{ animation: blink 1.0s infinite; }

.status-box{
  border: 1px solid #bfdbfe;
  background: #f8fbff;
  border-radius: 12px;
  padding: 10px 12px;
  min-height: 120px;
}
.status-box.success{ border-color:#bbf7d0; background:#f0fdf4; }
.status-box.warn{ border-color:#fde68a; background:#fff7ed; }
.status-box.error{ border-color:#fecaca; background:#fff1f2; }

.status-title{ font-weight: 900; font-size: 13px; color:#111827; }
.status-main{ font-size: 13px; margin-top: 6px; color:#111827; font-weight: 800; }
.status-sub { font-size: 12px; margin-top: 4px; color:#6b7280; }

/* ステータスを少し下げる用 */
.status-spacer{ height: 14px; }
</style>
""",
    unsafe_allow_html=True
)

# --- 2. AI設定（ロジック固定） ---
if "GEMINI_API_KEY" in st.secrets:
    pal_genai.configure(api_key=st.secrets["GEMINI_API_KEY"])
    api_ready = True
else:
    api_ready = False

# --- 3. ユーティリティ & マスタ読み込み（ロジック固定） ---
def clean_ai_text(text: str) -> str:
    text = re.sub(r'[\*\#\`]', '', text)
    return text.strip()

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
ASSETS_DIR = os.path.join(BASE_DIR, "assets")

def load_meeting_master():
    master_path = os.path.join(ASSETS_DIR, "会議マスタ.docx")
    master_data = {}
    if os.path.exists(master_path):
        doc = Document(master_path)
        full_text = "\n".join([p.text for p in doc.paragraphs])
        blocks = re.split(r'会議名：', full_text)
        for block in blocks:
            if '出席者：' in block:
                parts = block.split('出席者：')
                name = parts[0].strip()
                attendees = parts[1].strip().replace('\n', ' ')
                master_data[name] = attendees
    return master_data

def get_docx_text(path: str) -> str:
    return "\n".join([p.text for p in Document(path).paragraphs]) if os.path.exists(path) else ""

# --- 3.1 リトライ（指数バックオフ：ロジック固定） ---
def is_retryable_error(e: Exception) -> bool:
    msg = str(e).lower()
    return (
        "500" in msg
        or "503" in msg
        or "429" in msg
        or "internal error" in msg
        or "unavailable" in msg
        or "resource exhausted" in msg
        or "quota" in msg
        or "rate" in msg
        or "timeout" in msg
    )

def generate_with_retry(model, inputs, max_retries: int = 5):
    last = None
    for i in range(max_retries):
        try:
            return model.generate_content(inputs)
        except Exception as e:
            last = e
            if not is_retryable_error(e):
                raise
            wait = min(2 ** i, 30)
            time.sleep(wait)
    raise last

def safe_extract_text_from_response(response) -> tuple[str, str]:
    finish_reason_str = "UNKNOWN"
    if not getattr(response, "candidates", None):
        return "", finish_reason_str

    cand = response.candidates[0]
    fr = getattr(cand, "finish_reason", None)
    finish_reason_str = getattr(fr, "name", None) or str(fr)

    content = getattr(cand, "content", None)
    parts = getattr(content, "parts", None) if content else None
    if not parts:
        return "", finish_reason_str

    text = "".join(getattr(p, "text", "") for p in parts if getattr(p, "text", None))
    return text, finish_reason_str

MEETING_MASTER = load_meeting_master()
meeting_options = ["会議を選択してください"] + list(MEETING_MASTER.keys()) + ["その他"]

# --- 状態（ロジック固定） ---
if "safe_text" not in st.session_state:
    st.session_state.safe_text = ""
if "status_text" not in st.session_state:
    st.session_state.status_text = ""
if "status_level" not in st.session_state:
    st.session_state.status_level = "info"  # info / warning / error / success
if "finish_reason" not in st.session_state:
    st.session_state.finish_reason = ""
if "status_step_en" not in st.session_state:
    st.session_state.status_step_en = ""

# ===== UI（ロジック触らない） =====

# タイトル切れ最終対策：先頭スペーサ（UIのみ）
st.markdown("<div style='height:6px'></div>", unsafe_allow_html=True)

# ヘッダー
h_l, h_r = st.columns([7, 3], vertical_alignment="center")
with h_l:
    st.markdown("<div class='at-title'>Audio Transcription <span class='at-ver'>ver3.3</span></div>", unsafe_allow_html=True)
with h_r:
    st.markdown("<div class='model-chip'><div class='value'>Gemini-2.5-PRO</div></div>", unsafe_allow_html=True)

# 2列固定（5:5）＋中央スペース用の空カラム
left, gap, right = st.columns([5, 0.35, 5], vertical_alignment="top")
with gap:
    st.write("")  # 中央スペース

with left:
    selected_meeting = st.selectbox("会議名", meeting_options)

    p_l, p_r = st.columns([1, 1], vertical_alignment="top")
    with p_l:
        place = st.text_input("場所（Word反映）", value="本部会議室")
    with p_r:
        recorder = st.text_input("記録者（Word反映）", value="事務局")

    audio_files = st.file_uploader(
        "音声ファイル（複数可）",
        type=["mp3", "m4a", "wav"],
        accept_multiple_files=True
    )
    st.markdown("<div class='upload-help'>合計200MB以内（MP3 / M4A / WAV）</div>", unsafe_allow_html=True)

    st.markdown("<div class='section-spacer'></div>", unsafe_allow_html=True)
    extra_prompt = st.text_area(
        "追加指示（最優先）",
        placeholder="特記事項があれば入力してください",
        height=120
    )

    st.markdown("<div class='btn-spacer'></div>", unsafe_allow_html=True)
    start = st.button("🚀 生成を開始", use_container_width=True)

with right:
    mode = st.radio(
        "作成モード",
        ["議事録（フォーマット適用）", "発言録（全文記録）", "単純要約（セミナー等）", "箇条書き（報告者別）"],
        horizontal=True,
    )

    # How to use（薄いオレンジ + 番号位置調整）
    st.markdown(
        """
        <div class="howto-box">
          <div class="howto-title">How to use</div>
          <ol class="howto">
            <li>会議名選択（無い場合は「その他」）</li>
            <li>作成モード選択</li>
            <li>場所、記録者を入力</li>
            <li>音声ファイルをアップロード</li>
            <li>追加指示（あれば）を入力（最優先）</li>
            <li>「生成を開始」ボタンを押す</li>
            <li>ステータスが「生成完了」になったらプレビューで確認</li>
            <li>最下部の「Wordファイルを出力」を押す</li>
          </ol>
          <small>※ 出力とWordは一致します（v3.3）</small>
        </div>
        """,
        unsafe_allow_html=True
    )

    # ステータスを“少し下げる”
    st.markdown("<div class='status-spacer'></div>", unsafe_allow_html=True)

    level = st.session_state.status_level
    cls = "status-box"
    if level == "success":
        cls += " success"
    elif level == "warning":
        cls += " warn"
    elif level == "error":
        cls += " error"

    status_line = st.session_state.status_text or "待機中"
    step_en = st.session_state.status_step_en or ""
    finish = st.session_state.finish_reason
    finish_line = f"finish_reason: {finish}" if finish else ""

    blinking = ""
    if level == "info" and any(k in status_line for k in ["解析", "音声", "生成", "アップロード", "処理中"]):
        blinking = "blinking"

    st.markdown(
        f"""
        <div class="{cls} {blinking}">
          <div class="status-title">解析ステータス</div>
          <div class="status-main">{status_line}</div>
          <div class="status-sub">{step_en}</div>
          <div class="status-sub">{finish_line}</div>
        </div>
        """,
        unsafe_allow_html=True
    )

# 出席者マスタ（ロジック固定）
if selected_meeting == "会議を選択してください":
    attendee_master = ""
elif selected_meeting == "その他":
    attendee_master = "出席者不明"
else:
    attendee_master = MEETING_MASTER.get(selected_meeting, "")

# ===== 6. 解析処理（ロジックはそのまま） =====
if start:
    if not (api_ready and audio_files and selected_meeting != "会議を選択してください"):
        st.session_state.status_text = "設定と音声ファイルを確認してください。"
        st.session_state.status_level = "error"
        st.session_state.status_step_en = "Validation failed."
        st.rerun()
    else:
        st.session_state.status_text = "解析開始：音声をアップロードしています…"
        st.session_state.status_level = "info"
        st.session_state.status_step_en = "Uploading audio..."
        st.session_state.finish_reason = ""
        st.rerun()

if st.session_state.status_text.startswith("解析開始：") and st.session_state.status_level == "info":
    tmp_paths = []
    try:
        with st.status("🛰️ 解析中…", expanded=False):
            all_audio_data = []
            for idx, a_file in enumerate(audio_files, start=1):
                st.session_state.status_text = f"音声処理中…（{idx}/{len(audio_files)}）"
                st.session_state.status_level = "info"
                st.session_state.status_step_en = "Processing audio..."

                with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(a_file.name)[1]) as tmp:
                    tmp.write(a_file.getvalue())
                    tmp_path = tmp.name
                tmp_paths.append(tmp_path)

                audio_data = pal_genai.upload_file(path=tmp_path, mime_type="audio/mp4")
                
                while audio_data.state.name == "PROCESSING":
                    time.sleep(3)
                    audio_data = pal_genai.get_file(audio_data.name)

                if audio_data.state.name == "FAILED":
                    st.session_state.status_text = f"❌ ファイル処理に失敗しました: {a_file.name}"
                    st.session_state.status_level = "error"
                    st.session_state.status_step_en = "Audio processing failed."
                    st.stop()

                all_audio_data.append(audio_data)

            if "単純要約" in mode:
                p_filename = "prompt_simple_summary.docx"
            elif "箇条書き" in mode:
                p_filename = "prompt_bullet_reports.docx"
            else:
                p_filename = "prompt_gijiroku.docx" if "議事録" in mode else "prompt_hatsugen.docx"

            p_path = os.path.join(ASSETS_DIR, p_filename)
            final_prompt_base = get_docx_text(p_path)

            if ("単純要約" in mode) or ("箇条書き" in mode):
                extraction_rule = f"""
# 入力情報:
対象: {selected_meeting}
（出席者情報があれば参照）: {attendee_master}
"""
            else:
                extraction_rule = f"""
# 抽出の厳格ルール:
1. 音声データを解析し、「実際に発言（報告）が確認できた人物」のみを抽出し、その内容を記述せよ。
2. 文章の先頭に記号（*や-）を付けない。
3. 外部知識を排除し、事実のみを記述すること。
4. 特に【曲田課長】の報告末尾（AIアプリ開発、棚卸スケジュール等）を詳細に記述すること。

# 会議情報:
出席者: {attendee_master}
"""

            final_prompt = f"{final_prompt_base}\n\n{extraction_rule}\n\n# 会議名: {selected_meeting}\n# 追加指示: {extra_prompt}"

            safety_settings = [
                {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_NONE"},
                {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_NONE"},
                {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_NONE"},
                {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_NONE"},
            ]

            model = pal_genai.GenerativeModel(
                model_name="gemini-2.5-pro",
                generation_config={"temperature": 0},
                safety_settings=safety_settings,
            )

            st.session_state.status_text = "生成中…（Geminiがまとめています）"
            st.session_state.status_level = "info"
            st.session_state.status_step_en = "Generating transcript..."

            response = generate_with_retry(model, [final_prompt] + all_audio_data, max_retries=5)
            raw_text, finish_reason_str = safe_extract_text_from_response(response)

            st.session_state.finish_reason = finish_reason_str

            if raw_text.strip():
                st.session_state.safe_text = clean_ai_text(raw_text)
                st.session_state.status_text = "✅ 生成完了"
                st.session_state.status_level = "success"
                st.session_state.status_step_en = "Completed."
            else:
                st.session_state.status_text = "❌ 本文が返りませんでした。"
                st.session_state.status_level = "error"
                st.session_state.status_step_en = "No content returned."

    except Exception as e:
        st.session_state.status_text = f"❌ 解析エラー: {e}"
        st.session_state.status_level = "error"
        st.session_state.status_step_en = "Unexpected error."
    finally:
        for p in tmp_paths:
            try:
                if os.path.exists(p):
                    os.remove(p)
            except Exception:
                pass

    st.rerun()

# ===== 7. プレビュー & Word出力（ロジックそのまま） =====
if st.session_state.safe_text:
    st.markdown("<div style='border:1px solid #e5e7eb;background:#fff;border-radius:14px;padding:12px 12px;'>", unsafe_allow_html=True)
    st.subheader("📝 プレビュー（編集可）")
    st.session_state.safe_text = st.text_area(
        "生成結果（編集可）",
        st.session_state.safe_text,
        height=420,
        label_visibility="collapsed"
    )
    st.markdown("</div>", unsafe_allow_html=True)

    col_dl_l, col_dl_r = st.columns([2, 1], vertical_alignment="center")
    with col_dl_l:
        st.caption("※ ステータスが「生成完了」になったら内容確認 → Word出力")
    with col_dl_r:
        if st.button("📄 Wordファイルを出力", use_container_width=True):
            txt = st.session_state.safe_text
            today_str = datetime.now().strftime("%Y%m%d")
            display_date = datetime.now().strftime("%Y年%m月%d日")

            def replace_in_paragraph(paragraph, replace_map):
                full_text = paragraph.text
                for k, v in replace_map.items():
                    full_text = full_text.replace(k, str(v))
                paragraph.text = full_text

            def extract_section(tag_name, text):
                tag_match = re.search(rf"<{re.escape(tag_name)}>(.*?)</{re.escape(tag_name)}>", text, re.DOTALL)
                if tag_match:
                    return tag_match.group(1).strip()
                pattern = rf"(?:^|\n)[■\s]*{re.escape(tag_name)}.*?\n(.*?)(?=\n■|\n\d+\.\s|$)"
                match = re.search(pattern, text, re.DOTALL)
                return match.group(1).strip() if match else ""

            if "単純要約" in mode:
                template_file = "要約フォーマット.docx"
            else:
                template_file = "議事録フォーマット.docx" if ("議事録" in mode or "箇条書き" in mode) else "発言録フォーマット.docx"

            f_path = os.path.join(ASSETS_DIR, template_file)

            if not os.path.exists(f_path) and "単純要約" in mode:
                doc = Document()
                doc.add_paragraph(f"作成日：{display_date}")
                doc.add_paragraph(selected_meeting)
                doc.add_paragraph("")
                doc.add_paragraph("要約")
                doc.add_paragraph(txt)
                out_io = io.BytesIO()
                doc.save(out_io)
                out_io.seek(0)
                filename = f"{selected_meeting}_要約_{today_str}.docx"
                st.download_button("📥 ダウンロード", data=out_io, file_name=filename, use_container_width=True)
            else:
                if not os.path.exists(f_path):
                    st.error("❌ テンプレートファイルが見つかりませんでした。assetsフォルダを確認してください。")
                else:
                    doc = Document(f_path)

                    replace_map = {
                        "[[CREATED_DATE]]": display_date,
                        "[[MEETING_DATE]]": display_date,
                        "[[TITLE]]": selected_meeting,
                        "[[ATTENDEES]]": attendee_master,
                        "[[PLACE]]": place,
                        "[[RECORDER]]": recorder,
                    }

                    if "単純要約" in mode:
                        replace_map.update({"[[SUMMARY]]": txt})
                    elif "議事録" in mode or "箇条書き" in mode:
                        replace_map.update({
                            "[[OVERVIEW]]": extract_section("会議概要", txt),
                            "[[REPORTS_BLOCK]]": extract_section("各報告の詳細", txt) or extract_section("各報告", txt),
                            "[[EXEC_COMMENTS]]": extract_section("経営層コメント", txt),
                            "[[DECISIONS_BLOCK]]": extract_section("決定事項", txt),
                            "[[NEXT_MEETING]]": extract_section("次回予定", txt),
                        })
                    else:
                        replace_map.update({
                            "[[OVERVIEW]]": extract_section("T6", txt),
                            "[[LOG]]": extract_section("T7", txt) or txt,
                        })

                    for p in doc.paragraphs:
                        replace_in_paragraph(p, replace_map)

                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for p in cell.paragraphs:
                                    replace_in_paragraph(p, replace_map)
                                cell.text = cell.text.replace("*", "").strip()

                    out_io = io.BytesIO()
                    doc.save(out_io)
                    out_io.seek(0)

                    target_mode = "議事録" if "議事録" in mode else ("箇条書き" if "箇条書き" in mode else ("要約" if "単純要約" in mode else "発言録"))
                    filename = f"{selected_meeting}_{target_mode}_{today_str}.docx"
                    st.download_button("📥 ダウンロード", data=out_io, file_name=filename, use_container_width=True)